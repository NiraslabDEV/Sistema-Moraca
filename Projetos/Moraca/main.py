import ttkbootstrap as ttk
from ttkbootstrap.constants import *
import sqlite3
from datetime import datetime
import os
from tkinter import messagebox
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tkcalendar import DateEntry
import shutil
from PIL import Image, ImageDraw, ImageFont
import io
import base64
from datetime import datetime, timedelta
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
from openpyxl.drawing.image import Image
from os_interna_andamento import criar_os_interna_andamento, criar_os_ziehm
import subprocess
import platform
import sys
from licenca import verificar_licenca, carregar_licenca, gerar_nova_licenca_demo
from ativar_licenca import AtivadorLicenca
import re
from os_visita_tecnica import criar_os_visita_tecnica
from os_interna import criar_os_interna

# Classe de diálogo personalizada para substituir messagebox
class DialogoPersonalizado:
    def __init__(self, parent, titulo, mensagem, tipo="info"):
        self.resultado = None
        
        # Criar janela de diálogo
        self.janela = ttk.Toplevel(parent)
        self.janela.title(titulo)
        self.janela.geometry("400x200")
        self.janela.resizable(False, False)
        self.janela.transient(parent)  # Definir janela como transitória (dependente da principal)
        self.janela.grab_set()  # Torna o diálogo modal
        
        # Centralizar na tela
        self.janela.update_idletasks()
        largura = self.janela.winfo_width()
        altura = self.janela.winfo_height()
        x = (self.janela.winfo_screenwidth() // 2) - (largura // 2)
        y = (self.janela.winfo_screenheight() // 2) - (altura // 2)
        self.janela.geometry(f"{largura}x{altura}+{x}+{y}")
        
        # Frame principal
        frame = ttk.Frame(self.janela, padding=20)
        frame.pack(fill=ttk.BOTH, expand=ttk.YES)
        
        # Ícone baseado no tipo
        estilo_botao = "primary.TButton"
        if tipo == "erro":
            estilo_botao = "danger.TButton"
        elif tipo == "aviso":
            estilo_botao = "warning.TButton"
        elif tipo == "sucesso":
            estilo_botao = "success.TButton"
        
        # Mensagem
        ttk.Label(
            frame,
            text=mensagem,
            font=("Helvetica", 11),
            wraplength=360,
            justify=ttk.CENTER
        ).pack(pady=(10, 30))
        
        # Botão OK
        botao_ok = ttk.Button(
            frame,
            text="OK",
            command=self.confirmar,
            style=estilo_botao,
            width=15,
            cursor="hand2"  # Altera o cursor ao passar sobre o botão
        )
        botao_ok.pack(pady=10)
        
        # Configura tamanho maior para o botão
        botao_ok.configure(padding=(20, 10))
        
        # Definir foco no botão OK
        botao_ok.focus_set()
        
        # Vincular tecla Enter para confirmar
        self.janela.bind("<Return>", lambda event: self.confirmar())
        self.janela.bind("<Escape>", lambda event: self.confirmar())
        
        # Aguardar interação do usuário
        parent.wait_window(self.janela)
    
    def confirmar(self):
        self.resultado = True
        self.janela.destroy()

# Função para substituir messagebox.showinfo
def mostrar_mensagem(parent, titulo, mensagem, tipo="info"):
    dialogo = DialogoPersonalizado(parent, titulo, mensagem, tipo)
    return dialogo.resultado

# Importar o módulo para gerar a OS interna de andamento
try:
    from os_interna_andamento import criar_os_interna_andamento
except ImportError:
    print("Módulo os_interna_andamento não encontrado, algumas funcionalidades estarão indisponíveis")

class MoracaOS:
    def __init__(self):
        # Verificar a licença antes de continuar
        if not self.verificar_licenca_sistema():
            return
            
        # Definir cores personalizadas para a interface
        self.cores = {
            # Cores para status das OS
            "aberta": "#005452",    # Azul esverdeado
            "fechada": "#515151",   # Verde escuro
            
            # Cor para resultados de busca
            "encontrado": "#ccf2ff",  # Azul mais claro
            
            # Cores para os botões
            "botao_primario": "primary",
            "botao_secundario": "secondary",
            "botao_info": "info",
            "botao_sucesso": "success",
            "botao_perigo": "danger",
            "botao_aviso": "warning"
        }
        
        # Mapeamento dos status antigos para os novos (para compatibilidade)
        self.mapeamento_status = {
            "Pendente": "ABERTA",
            "Em Andamento": "ABERTA",
            "Concluída": "FECHADA",
        }
        
        self.root = ttk.Window(themename="solar")
        self.root.title("Moraca - Sistema de OS")
        # Aumentar o tamanho inicial da janela
        self.root.geometry("900x700")
        
        # Criar estrutura de pastas
        self.criar_estrutura_pastas()
        
        # Verificar e copiar logo se necessário
        self.verificar_logo()
        
        # Criar banco de dados
        self.criar_banco()
        
        # Container principal para permitir scrolling
        main_container = ttk.Frame(self.root)
        main_container.pack(fill=BOTH, expand=YES)
        
        # Canvas para aplicar scrolling na tela inteira
        self.main_canvas = ttk.Canvas(main_container)
        self.main_canvas.pack(side=LEFT, fill=BOTH, expand=YES)
        
        # Scrollbar vertical
        main_scrollbar = ttk.Scrollbar(main_container, orient=VERTICAL, command=self.main_canvas.yview)
        main_scrollbar.pack(side=RIGHT, fill=Y)
        self.main_canvas.configure(yscrollcommand=main_scrollbar.set)
        
        # Frame principal que irá conter todos os elementos
        self.main_frame = ttk.Frame(self.main_canvas, padding="20")
        
        # Criar janela dentro do canvas
        self.frame_window = self.main_canvas.create_window((0, 0), window=self.main_frame, anchor=NW)
        
        # Configurar eventos para redimensionamento
        self.main_frame.bind("<Configure>", self._configure_main_frame)
        self.main_canvas.bind("<Configure>", self._configure_main_canvas)
        
        # Configurar eventos de scroll
        self._configure_mousewheel_scrolling()
        
        # Título
        ttk.Label(
            self.main_frame,
            text="Moraca - Sistema de Ordem de Serviço",
            font=("Helvetica", 16, "bold")
        ).pack(pady=10)
        
        # Frame para logos
        self.logos_frame = ttk.Frame(self.main_frame)
        self.logos_frame.pack(pady=10)
        
        # Adicionar os logos
        self.carregar_logos(self.logos_frame)
        
        # Frame para pesquisa (reposicionado para estar logo após o frame de logos)
        self.search_frame = ttk.Frame(self.main_frame)
        self.search_frame.pack(pady=10)
        
        ttk.Label(
            self.search_frame,
            text="Pesquisar OS:"
        ).pack(side=LEFT, padx=5)
        
        self.search_entry = ttk.Entry(self.search_frame, width=30)
        self.search_entry.pack(side=LEFT, padx=5)
        
        ttk.Button(
            self.search_frame,
            text="Buscar",
            command=self.buscar_os,
            style="secondary.TButton"
        ).pack(side=LEFT, padx=5)
        
        ttk.Button(
            self.search_frame,
            text="Ver Todas",
            command=self.mostrar_todas_os,
            style="secondary.Outline.TButton"
        ).pack(side=LEFT, padx=5)
        
        # Frame para botões
        self.button_frame = ttk.Frame(self.main_frame)
        self.button_frame.pack(pady=10)
        
        # Botão principal
        ttk.Button(
            self.button_frame,
            text="Nova OS",
            command=self.criar_nova_os,
            style="primary.TButton",
            width=30
        ).pack(pady=10)
        
        # Botão para gerar OS interna de andamento
        ttk.Button(
            self.button_frame,
            text="Gerar OS Interna de Andamento",
            command=self.gerar_os_interna_andamento,
            style="secondary.TButton",
            width=30
        ).pack(pady=10)
        
        # Separador para botões de impressão
        ttk.Label(
            self.button_frame,
            text="Impressão Rápida (selecione uma OS na lista)",
            font=("Helvetica", 10, "bold")
        ).pack(pady=(10, 5))
        
        # Frame para primeira linha de botões de impressão
        print_buttons_frame1 = ttk.Frame(self.button_frame)
        print_buttons_frame1.pack(pady=5)
        
        # Botões de impressão primeira linha
        ttk.Button(
            print_buttons_frame1,
            text="OS ANDAMENTO",
            command=self.imprimir_os_andamento_rapida,
            style="info.TButton", 
            width=15
        ).pack(side=LEFT, padx=5)
        
        ttk.Button(
            print_buttons_frame1,
            text="OS INTERNA",
            command=self.imprimir_os_interna_rapida,
            style="info.TButton",
            width=15
        ).pack(side=LEFT, padx=5)
        
        ttk.Button(
            print_buttons_frame1,
            text="OS VISITA ZIEHM",
            command=self.imprimir_os_ziehm_rapida,
            style="info.TButton",
            width=15
        ).pack(side=LEFT, padx=5)
        
        # Frame para segunda linha de botões de impressão
        print_buttons_frame2 = ttk.Frame(self.button_frame)
        print_buttons_frame2.pack(pady=5)
        
        # Botões de impressão segunda linha
        ttk.Button(
            print_buttons_frame2,
            text="OS PREVENTIVA",
            command=self.imprimir_os_preventiva_rapida,
            style="info.TButton",
            width=15
        ).pack(side=LEFT, padx=5)
        
        ttk.Button(
            print_buttons_frame2,
            text="OS VISITA CORRETIVA",
            command=self.imprimir_os_visita_corretiva_rapida,
            style="info.TButton",
            width=15
        ).pack(side=LEFT, padx=5)
        
        ttk.Button(
            print_buttons_frame2,
            text="OS VISITA TÉCNICA",
            command=self.imprimir_os_visita_tecnica_rapida,
            style="info.TButton",
            width=15
        ).pack(side=LEFT, padx=5)
        
        # Lista de últimas OS
        self.lista_frame = ttk.Frame(self.main_frame)
        # Modificar para expandir mais e dar mais espaço à lista
        self.lista_frame.pack(fill=BOTH, expand=YES, pady=10)
        
        ttk.Label(
            self.lista_frame,
            text="Últimas OS Criadas",
            font=("Helvetica", 12, "bold")
        ).pack(pady=5)  # Reduzir padding para dar mais espaço à tabela
        
        # Criar frame para conter a tabela e as barras de rolagem
        tabela_frame = ttk.Frame(self.lista_frame)
        tabela_frame.pack(fill=BOTH, expand=YES, padx=5, pady=5)
        
        # Criar barra de rolagem vertical
        y_scrollbar = ttk.Scrollbar(tabela_frame, orient=VERTICAL)
        y_scrollbar.pack(side=RIGHT, fill=Y)
        
        # Criar barra de rolagem horizontal
        x_scrollbar = ttk.Scrollbar(tabela_frame, orient=HORIZONTAL)
        x_scrollbar.pack(side=BOTTOM, fill=X)
        
        # Criar tabela
        self.tabela = ttk.Treeview(
            tabela_frame,
            columns=("numero", "cliente", "data", "status"),
            show="headings",
            yscrollcommand=y_scrollbar.set,
            xscrollcommand=x_scrollbar.set,
            height=15  # Aumentar altura da tabela para mostrar mais linhas
        )
        
        # Configurar barras de rolagem para controlar a tabela
        y_scrollbar.config(command=self.tabela.yview)
        x_scrollbar.config(command=self.tabela.xview)
        
        # Configurar cabeçalhos e larguras das colunas
        self.tabela.heading("numero", text="Número")
        self.tabela.heading("cliente", text="Cliente")
        self.tabela.heading("data", text="Data")
        self.tabela.heading("status", text="Status")
        
        # Definir larguras para as colunas
        self.tabela.column("numero", width=120, minwidth=100)
        self.tabela.column("cliente", width=300, minwidth=200)
        self.tabela.column("data", width=100, minwidth=80)
        self.tabela.column("status", width=150, minwidth=100)
        
        self.tabela.pack(fill=BOTH, expand=YES)
        
        # Adicionar evento de clique duplo para abrir a aba de impressão
        self.tabela.bind("<Double-1>", self.abrir_aba_impressao)
        
        # Atualizar lista
        self.atualizar_lista()
    
    def criar_estrutura_pastas(self):
        # Caminho base dos documentos técnicos
        base_path = os.path.join("MORACA", "MORACA1", "DOCUMENTOS", "tecnico")
        if not os.path.exists(base_path):
            os.makedirs(base_path)
            
        # Estrutura anterior (manter para compatibilidade)
        for pasta in ["andamento", "preventivas", "visitas"]:
            pasta_path = os.path.join(base_path, pasta)
            if not os.path.exists(pasta_path):
                os.makedirs(pasta_path)
        
        # Criar pasta principal se não existir
        if not os.path.exists("OSs"):
            os.makedirs("OSs")
        
        # Criar subpastas se não existirem
        subpastas = ["Pendentes", "Em_Andamento", "Concluidas"]
        for subpasta in subpastas:
            if not os.path.exists(os.path.join("OSs", subpasta)):
                os.makedirs(os.path.join("OSs", subpasta))
                
        # Nova estrutura de pastas por ano e mês
        self.criar_estrutura_os_ano_mes()
    
    def criar_estrutura_os_ano_mes(self):
        """Cria a estrutura de pastas organizadas por ano e mês para as OS"""
        # Caminho base para a nova organização
        base_path = os.path.join("MORACA", "MORACA1", "DOCUMENTOS", "tecnico", "os")
        if not os.path.exists(base_path):
            os.makedirs(base_path)
            
        # Criar pasta do ano atual
        ano_atual = datetime.now().strftime("%Y")
        pasta_ano = os.path.join(base_path, f"OS-{ano_atual}")
        if not os.path.exists(pasta_ano):
            os.makedirs(pasta_ano)
            
        # Criar pasta do mês atual
        meses = {
            "01": "janeiro",
            "02": "fevereiro",
            "03": "marco",
            "04": "abril",
            "05": "maio",
            "06": "junho",
            "07": "julho",
            "08": "agosto",
            "09": "setembro",
            "10": "outubro",
            "11": "novembro",
            "12": "dezembro"
        }
        
        mes_atual = datetime.now().strftime("%m")
        nome_mes = meses.get(mes_atual, "outros")
        pasta_mes = os.path.join(pasta_ano, f"os-{nome_mes}")
        if not os.path.exists(pasta_mes):
            os.makedirs(pasta_mes)
    
    def get_pasta_os_cliente(self, numero, cliente):
        """
        Obtém o caminho da pasta específica para uma OS
        Formato: /MORACA/MORACA1/DOCUMENTOS/tecnico/os/OS-[ANO]/os-[MES]/OS [ANO] [NUMERO] - [CLIENTE]
        
        Args:
            numero: Número da OS (ex: OS2519)
            cliente: Nome do cliente
            
        Returns:
            str: Caminho completo da pasta para a OS específica
        """
        # Extrair ano e número da OS
        match = re.search(r'OS(\d{2})(\d+)', numero)
        if match:
            ano = match.group(1)
            sequencia = match.group(2)
            ano_completo = f"20{ano}"
        else:
            # Se não conseguir extrair, usar ano atual
            ano = datetime.now().strftime("%y")
            sequencia = re.sub(r'[^\d]', '', numero)
            ano_completo = datetime.now().strftime("%Y")
        
        # Determinar o mês baseado na data de criação da OS
        # Para este exemplo, usaremos o mês atual, mas idealmente seria baseado na data da OS
        meses = {
            "01": "janeiro",
            "02": "fevereiro",
            "03": "marco",
            "04": "abril",
            "05": "maio",
            "06": "junho",
            "07": "julho",
            "08": "agosto",
            "09": "setembro",
            "10": "outubro",
            "11": "novembro",
            "12": "dezembro"
        }
        
        mes_atual = datetime.now().strftime("%m")
        nome_mes = meses.get(mes_atual, "outros")
        
        # Nome da pasta específica da OS
        nome_cliente_simplificado = re.sub(r'[^\w\s-]', '', cliente).strip()[:30]  # Limitar tamanho e remover caracteres especiais
        nome_pasta_os = f"OS {ano} {sequencia} - {nome_cliente_simplificado}"
        
        # Caminho completo
        base_path = os.path.join("MORACA", "MORACA1", "DOCUMENTOS", "tecnico", "os")
        pasta_ano = os.path.join(base_path, f"OS-{ano_completo}")
        pasta_mes = os.path.join(pasta_ano, f"os-{nome_mes}")
        pasta_os = os.path.join(pasta_mes, nome_pasta_os)
        
        # Garantir que as pastas existam
        if not os.path.exists(pasta_ano):
            os.makedirs(pasta_ano)
        if not os.path.exists(pasta_mes):
            os.makedirs(pasta_mes)
        if not os.path.exists(pasta_os):
            os.makedirs(pasta_os)
            
        return pasta_os
    
    def verificar_logo(self):
        """Verifica se a logo existe na pasta assets, se não existir, cria uma logo placeholder"""
        if not os.path.exists("assets"):
            os.makedirs("assets")
            
        logo_path = os.path.join("assets", "logo.png")
        if not os.path.exists(logo_path):
            try:
                # Criar uma logo placeholder
                from PIL import Image, ImageDraw, ImageFont
                
                # Criar uma imagem com fundo branco
                img = Image.new('RGB', (400, 100), color=(255, 255, 255))
                draw = ImageDraw.Draw(img)
                
                # Desenhar texto
                draw.text((10, 40), "MORACA SISTEMAS", fill=(0, 0, 0))
                
                # Salvar a imagem
                img.save(logo_path)
                print("Foi criada uma imagem placeholder em assets/logo.png")
                print("Substitua essa imagem pela logo real antes de usar o sistema")
            except Exception as e:
                print(f"Erro ao criar logo placeholder: {e}")
                print("Arquivo logo.png não encontrado!")
                print("Por favor, salve a imagem da logo em assets/logo.png")
                print("A imagem deve ter o formato PNG e resolução adequada")
                print("Após salvar a imagem no local correto, execute o programa novamente")
    
    def carregar_logos(self, frame):
        """Carrega e exibe os logos da Moraca e Ziehm"""
        try:
            from PIL import Image, ImageTk
            
            # Criar um frame central para conter os dois logos
            logos_central_frame = ttk.Frame(frame)
            logos_central_frame.pack(pady=5)
            
            # Logo da Moraca
            logo_moraca_path = os.path.join("assets", "logo.png")
            if os.path.exists(logo_moraca_path):
                # Criar um frame para o logo da Moraca com borda estilizada
                moraca_frame = ttk.Frame(logos_central_frame, style="TFrame", borderwidth=2, relief="groove")
                moraca_frame.pack(side=LEFT, padx=25, pady=5)
                
                # Abrir e redimensionar a imagem
                img_moraca = Image.open(logo_moraca_path)
                # Calcular proporcionalidade
                width, height = img_moraca.size
                new_width = 180
                new_height = int(height * (new_width / width))
                img_moraca = img_moraca.resize((new_width, new_height), Image.LANCZOS)
                
                # Converter para PhotoImage
                logo_moraca = ImageTk.PhotoImage(img_moraca)
                
                # Manter uma referência para evitar garbage collection
                self.logo_moraca_img = logo_moraca
                
                # Adicionar o logo em um label com padding
                logo_label = ttk.Label(moraca_frame, image=logo_moraca)
                logo_label.pack(padx=10, pady=10)
                ttk.Label(moraca_frame, text="MORACA SISTEMAS", font=("Helvetica", 10, "bold")).pack(pady=(0, 10))
            
            # Separador vertical
            ttk.Separator(logos_central_frame, orient=VERTICAL).pack(side=LEFT, fill=Y, padx=15, pady=10)
            
            # Logo da Ziehm
            ziehm_logo_files = ["ziehmlogo.png", "logoziehm.jpg"]
            ziehm_logo_path = None
            
            for logo_file in ziehm_logo_files:
                temp_path = os.path.join("assets", logo_file)
                if os.path.exists(temp_path):
                    ziehm_logo_path = temp_path
                    break
            
            if ziehm_logo_path:
                # Criar um frame para o logo da Ziehm com borda estilizada
                ziehm_frame = ttk.Frame(logos_central_frame, style="TFrame", borderwidth=2, relief="groove")
                ziehm_frame.pack(side=LEFT, padx=25, pady=5)
                
                # Abrir e redimensionar a imagem
                img_ziehm = Image.open(ziehm_logo_path)
                # Calcular proporcionalidade
                width, height = img_ziehm.size
                new_width = 180
                new_height = int(height * (new_width / width))
                img_ziehm = img_ziehm.resize((new_width, new_height), Image.LANCZOS)
                
                # Converter para PhotoImage
                logo_ziehm = ImageTk.PhotoImage(img_ziehm)
                
                # Manter uma referência para evitar garbage collection
                self.logo_ziehm_img = logo_ziehm
                
                # Adicionar o logo em um label com padding
                logo_label = ttk.Label(ziehm_frame, image=logo_ziehm)
                logo_label.pack(padx=10, pady=10)
                ttk.Label(ziehm_frame, text="ZIEHM IMAGING", font=("Helvetica", 10, "bold")).pack(pady=(0, 10))
                
        except Exception as e:
            print(f"Erro ao carregar logos: {e}")
            # Criar um label com texto em vez das imagens
            ttk.Label(frame, text="MORACA SISTEMAS - ZIEHM IMAGING", font=("Helvetica", 12, "bold")).pack()
    
    def criar_banco(self):
        conn = sqlite3.connect('moraca.db')
        c = conn.cursor()
        c.execute('''CREATE TABLE IF NOT EXISTS os
                    (numero TEXT PRIMARY KEY,
                     cliente TEXT,
                     maquina TEXT,
                     descricao TEXT,
                     urgencia TEXT,
                     data TEXT,
                     status TEXT,
                     tipo TEXT,
                     patrimonio TEXT,
                     numero_serie TEXT,
                     local TEXT,
                     endereco TEXT,
                     cidade TEXT,
                     telefone TEXT,
                     cep TEXT,
                     contato_nome TEXT,
                     contato_telefone1 TEXT,
                     contato_telefone2 TEXT,
                     descricao_servico TEXT,
                     necessita_viagem TEXT,
                     tipo_hospedagem TEXT,
                     prazo_entrega TEXT,
                     empresa TEXT)''')
        conn.commit()
        
        # Verificar se a coluna 'empresa' existe, e adicionar se não existir
        c.execute("PRAGMA table_info(os)")
        colunas = [info[1] for info in c.fetchall()]
        if 'empresa' not in colunas:
            try:
                c.execute("ALTER TABLE os ADD COLUMN empresa TEXT")
                print("Coluna 'empresa' adicionada com sucesso ao banco de dados.")
                conn.commit()
            except sqlite3.OperationalError as e:
                print(f"Erro ao adicionar coluna 'empresa': {e}")
        
        conn.close()
    
    def gerar_arquivo_excel_os(self, numero, cliente, maquina, descricao, urgencia, data, status, tipo,
                        patrimonio, numero_serie, local, endereco, cidade, telefone, cep,
                        contato_nome, contato_telefone1, contato_telefone2, descricao_servico,
                        necessita_viagem, tipo_hospedagem, prazo_entrega, empresa):
        # Criar uma nova planilha
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = f"OS {numero}"
        
        # Ajustar a altura da linha
        ws.row_dimensions[1].height = 25
        
        # Configurar larguras das colunas
        ws.column_dimensions['A'].width = 15
        ws.column_dimensions['B'].width = 15
        ws.column_dimensions['C'].width = 15  # Ajustei para ter 3 colunas iguais
        ws.column_dimensions['D'].width = 30  # Aumentada para a célula do prazo de entrega
        ws.column_dimensions['E'].width = 20  # Aumentada para a célula do prazo de entrega
        
        # Ajustar a altura da linha do prazo
        ws.row_dimensions[3].height = 30  # Altura aumentada para a linha do prazo
        
        # Definir estilos
        titulo_font = Font(name='Arial', size=16, bold=True)
        cabecalho_font = Font(name='Arial', size=11, bold=True)
        normal_font = Font(name='Arial', size=10)
        center_align = Alignment(horizontal='center', vertical='center')
        left_align = Alignment(horizontal='left', vertical='center')
        
        # Borda mais espessa para compatibilidade com LibreOffice
        medium_border = Border(
            left=Side(style='medium'),
            right=Side(style='medium'),
            top=Side(style='medium'),
            bottom=Side(style='medium')
        )
        
        # Título
        ws.merge_cells('A1:E1')
        ws['A1'] = f"O.S. {numero}"
        ws['A1'].font = titulo_font
        ws['A1'].alignment = center_align
        ws['A1'].border = medium_border
        
        # Cliente e Prazo
        ws.merge_cells('A2:E2')
        ws['A2'] = f"Cliente: {cliente}"
        ws['A2'].font = cabecalho_font
        ws['A2'].alignment = center_align
        ws['A2'].border = medium_border
        
        ws.merge_cells('A3:E3')
        ws['A3'] = f"Prazo de Entrega Final: {prazo_entrega}"
        ws['A3'].font = cabecalho_font
        ws['A3'].alignment = center_align
        ws['A3'].border = medium_border
        
        # Informações básicas
        info_rows = [
            ("Equipamento:", maquina),
            ("Número de Série:", numero_serie),
            ("Nº de Patrimônio:", patrimonio),
            ("Data de abertura O.S.I.:", data),
            ("Local:", local)
        ]
        
        row = 4
        for label, valor in info_rows:
            ws.merge_cells(f'A{row}:E{row}')
            ws[f'A{row}'].font = normal_font
            ws[f'A{row}'].alignment = center_align
            ws[f'A{row}'].border = medium_border
            ws[f'A{row}'] = f"{label} {valor}"
            row += 1
        
        # Tabela de tarefas
        ws.merge_cells(f'A{row}:C{row}')
        ws[f'A{row}'] = "TAREFA A EXECUTAR"
        ws[f'A{row}'].font = cabecalho_font
        ws[f'A{row}'].alignment = center_align
        ws[f'A{row}'].border = medium_border
        
        ws[f'D{row}'] = "PRAZO PREVISTO"
        ws[f'D{row}'].font = cabecalho_font
        ws[f'D{row}'].alignment = center_align
        ws[f'D{row}'].border = medium_border
        
        ws[f'E{row}'] = "OK"
        ws[f'E{row}'].font = cabecalho_font
        ws[f'E{row}'].alignment = center_align
        ws[f'E{row}'].border = medium_border
        row += 1
        
        # Linhas para tarefas (vazias)
        for i in range(10):
            ws.merge_cells(f'A{row}:C{row}')
            ws[f'A{row}'].border = medium_border
            ws[f'D{row}'].border = medium_border
            ws[f'E{row}'].border = medium_border
            row += 1
        
        # Descrição / Observações
        ws.merge_cells(f'A{row}:E{row}')
        ws[f'A{row}'] = "OBSERVAÇÕES:"
        ws[f'A{row}'].font = cabecalho_font
        ws[f'A{row}'].alignment = center_align
        ws[f'A{row}'].border = medium_border
        row += 1
        
        # Texto completo da descrição
        ws.merge_cells(f'A{row}:E{row+5}')
        ws[f'A{row}'] = descricao
        ws[f'A{row}'].font = normal_font
        ws[f'A{row}'].alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
        ws[f'A{row}'].border = medium_border
        
        # Configurações de impressão
        ws.print_area = 'A1:E30'
        ws.page_setup.orientation = 'portrait'
        ws.page_setup.paperSize = ws.PAPERSIZE_A4
        ws.page_setup.fitToPage = True
        ws.page_setup.fitToHeight = 1
        ws.page_setup.fitToWidth = 1
        
        # Obter o caminho da pasta específica para esta OS
        pasta_os = self.get_pasta_os_cliente(numero, cliente)
        
        # Salvar arquivo Excel na nova estrutura
        arquivo_excel = os.path.join(pasta_os, f"{numero}.xlsx")
        wb.save(arquivo_excel)
        
        print(f"Arquivo Excel salvo em: {arquivo_excel}")
        
        # Manter o salvamento original para compatibilidade
        pasta_status = {"ABERTA": "Pendentes", "FECHADA": "Concluidas"}.get(status, "Pendentes")
        pasta_compat = os.path.join("OSs", pasta_status, numero)
        if not os.path.exists(pasta_compat):
            os.makedirs(pasta_compat)
        
        arquivo_excel_compat = os.path.join(pasta_compat, f"{numero}.xlsx")
        wb.save(arquivo_excel_compat)
    
    def gerar_arquivo_os(self, numero, cliente, maquina, descricao, urgencia, data, status, tipo,
                        patrimonio, numero_serie, local, endereco, cidade, telefone, cep,
                        contato_nome, contato_telefone1, contato_telefone2, descricao_servico,
                        necessita_viagem, tipo_hospedagem, prazo_entrega, empresa):
        # Criar documento
        doc = Document()
        
        # Estilo para título
        title_style = doc.styles['Title']
        title_style.font.size = Pt(16)
        title_style.font.bold = True
        
        # Estilo para cabeçalho
        heading_style = doc.styles['Heading 1']
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True
        
        # Título do documento
        doc.add_heading(f"Ordem de Serviço: {numero}", 0)
        
        # Adicionar logo
        self.adicionar_cabecalho(doc, numero)
        
        # Adicionar informações da OS
        p = doc.add_paragraph()
        p.add_run("Empresa: ").bold = True
        p.add_run(empresa)
        
        p = doc.add_paragraph()
        p.add_run("Cliente: ").bold = True
        p.add_run(cliente)
        
        p = doc.add_paragraph()
        p.add_run("Máquina: ").bold = True
        p.add_run(maquina)
        
        # Adicionar campos novos
        if numero_serie:
            p = doc.add_paragraph()
            p.add_run("Número de Série: ").bold = True
            p.add_run(numero_serie)
            
        if patrimonio:
            p = doc.add_paragraph()
            p.add_run("Patrimônio: ").bold = True
            p.add_run(patrimonio)
            
        if local:
            p = doc.add_paragraph()
            p.add_run("Local: ").bold = True
            p.add_run(local)
            
        if endereco:
            p = doc.add_paragraph()
            p.add_run("Endereço: ").bold = True
            p.add_run(endereco)
            
        if cidade:
            p = doc.add_paragraph()
            p.add_run("Cidade: ").bold = True
            p.add_run(cidade)
            
        if telefone:
            p = doc.add_paragraph()
            p.add_run("Telefone: ").bold = True
            p.add_run(telefone)
            
        if contato_nome:
            p = doc.add_paragraph()
            p.add_run("Contato: ").bold = True
            p.add_run(contato_nome)
        
        if contato_telefone1:
            p = doc.add_paragraph()
            p.add_run("Telefone Contato: ").bold = True
            p.add_run(contato_telefone1)
            
        if contato_telefone2:
            p = doc.add_paragraph()
            p.add_run("Telefone Contato Alternativo: ").bold = True
            p.add_run(contato_telefone2)
            
        # Campos adicionais
        p = doc.add_paragraph()
        p.add_run("Urgência: ").bold = True
        p.add_run(urgencia)
        
        p = doc.add_paragraph()
        p.add_run("Tipo de OS: ").bold = True
        p.add_run(tipo)
        
        p = doc.add_paragraph()
        p.add_run("Data: ").bold = True
        p.add_run(data)
        
        if necessita_viagem:
            p = doc.add_paragraph()
            p.add_run("Necessita Viagem: ").bold = True
            p.add_run(necessita_viagem)
            
        if tipo_hospedagem:
            p = doc.add_paragraph()
            p.add_run("Tipo de Hospedagem: ").bold = True
            p.add_run(tipo_hospedagem)
            
        if prazo_entrega:
            p = doc.add_paragraph()
            p.add_run("Prazo de Entrega: ").bold = True
            p.add_run(prazo_entrega)
        
        # Descrição do problema
        doc.add_heading("Descrição do Problema", 1)
        doc.add_paragraph(descricao)
        
        # Descrição do serviço a ser executado
        if descricao_servico:
            doc.add_heading("Descrição do Serviço", 1)
            doc.add_paragraph(descricao_servico)
        
        # Obter o caminho da pasta específica para esta OS
        pasta_os = self.get_pasta_os_cliente(numero, cliente)
        
        # Garantir que a pasta existe
        if not os.path.exists(pasta_os):
            os.makedirs(pasta_os)
            
        # Salvar documento Word na nova estrutura
        arquivo_word = os.path.join(pasta_os, f"{numero}.docx")
        doc.save(arquivo_word)
        
        print(f"Arquivo Word salvo em: {arquivo_word}")
        
        # Manter o salvamento original para compatibilidade
        pasta_status = {"ABERTA": "Pendentes", "FECHADA": "Concluidas"}.get(status, "Pendentes")
        pasta_compat = os.path.join("OSs", pasta_status, numero)
        if not os.path.exists(pasta_compat):
            os.makedirs(pasta_compat)
            
        arquivo_word_compat = os.path.join(pasta_compat, f"{numero}.docx")
        doc.save(arquivo_word_compat)
    
    def criar_nova_os(self):
        # Criar nova janela
        nova_janela = ttk.Toplevel(self.root)
        nova_janela.title("Nova Ordem de Serviço")
        nova_janela.geometry("700x850")
        
        # Container principal
        main_container = ttk.Frame(nova_janela)
        main_container.pack(fill=BOTH, expand=TRUE)
        
        # Canvas para permitir scrolling
        canvas = ttk.Canvas(main_container)
        canvas.pack(side=LEFT, fill=BOTH, expand=TRUE)
        
        # Scrollbar
        scrollbar = ttk.Scrollbar(main_container, orient=VERTICAL, command=canvas.yview)
        scrollbar.pack(side=RIGHT, fill=Y)
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Frame interior para os elementos do formulário
        form_frame = ttk.Frame(canvas)
        
        # Criar janela dentro do canvas
        window_id = canvas.create_window((0, 0), window=form_frame, anchor=NW, width=680)
        
        # Função para atualizar scrollregion
        def _configure_frame(event):
            canvas.configure(scrollregion=canvas.bbox("all"))
        
        # Função para ajustar largura da janela interna
        def _configure_canvas(event):
            canvas.itemconfig(window_id, width=event.width)
        
        # Função de scroll multiplataforma
        def _on_mousewheel(event):
            # Linux: evento Button-4 (rolar para cima) e Button-5 (rolar para baixo)
            if event.num == 4:
                canvas.yview_scroll(-1, "units")
            elif event.num == 5:
                canvas.yview_scroll(1, "units")
            # Windows: evento MouseWheel com delta
            elif hasattr(event, 'delta'):
                canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        
        # Vincular eventos multiplataforma
        form_frame.bind("<Configure>", _configure_frame)
        canvas.bind("<Configure>", _configure_canvas)
        
        # Vincular eventos de scroll específicos por plataforma
        # Windows/macOS
        canvas.bind_all("<MouseWheel>", _on_mousewheel)
        # Linux
        canvas.bind_all("<Button-4>", _on_mousewheel)
        canvas.bind_all("<Button-5>", _on_mousewheel)
    
        # Organizando o formulário em seções
        current_row = 0
        
        # SEÇÃO DE SELEÇÃO DE EMPRESA (NOVA)
        ttk.Label(form_frame, text="SELEÇÃO DE EMPRESA", font=("Helvetica", 12, "bold")).grid(row=current_row, column=0, columnspan=2, sticky=W, padx=10, pady=(20,5))
        current_row += 1
        ttk.Separator(form_frame, orient=HORIZONTAL).grid(row=current_row, column=0, columnspan=2, sticky=EW, padx=10, pady=5)
        current_row += 1
        
        # Campo de Seleção de Empresa
        ttk.Label(form_frame, text="Empresa:", font=("Helvetica", 11, "bold")).grid(row=current_row, column=0, sticky=W, padx=10, pady=10)
        empresa_var = ttk.StringVar(value="")
        empresa_frame = ttk.Frame(form_frame)
        empresa_frame.grid(row=current_row, column=1, sticky=W, padx=10, pady=10)
        
        ttk.Radiobutton(empresa_frame, text="MORACA", value="MORACA", variable=empresa_var).pack(side=LEFT, padx=10)
        ttk.Radiobutton(empresa_frame, text="ZIEHM", value="ZIEHM", variable=empresa_var).pack(side=LEFT, padx=10)
        ttk.Radiobutton(empresa_frame, text="MORACA/ZIEHM", value="MORACA/ZIEHM", variable=empresa_var).pack(side=LEFT, padx=10)
        
        # Informações sobre campo obrigatório
        ttk.Label(form_frame, text="* Campo obrigatório", font=("Helvetica", 8), foreground="red").grid(row=current_row+1, column=1, sticky=W, padx=10)
        current_row += 2
        
        # SEÇÃO 1: INFORMAÇÕES BÁSICAS
        ttk.Label(form_frame, text="INFORMAÇÕES BÁSICAS", font=("Helvetica", 12, "bold")).grid(row=current_row, column=0, columnspan=2, sticky=W, padx=10, pady=(20,5))
        current_row += 1
        ttk.Separator(form_frame, orient=HORIZONTAL).grid(row=current_row, column=0, columnspan=2, sticky=EW, padx=10, pady=5)
        current_row += 1
        
        # Tipo de OS
        ttk.Label(form_frame, text="Tipo de OS:").grid(row=current_row, column=0, sticky=W, padx=10, pady=5)
        tipo_var = ttk.StringVar(value="MANUTENÇÃO CORRETIVA")
        tipo_combo = ttk.Combobox(
            form_frame,
            textvariable=tipo_var,
            values=["MANUTENÇÃO CORRETIVA", "MANUTENÇÃO PREVENTIVA", "ENTRADA DE MÁQUINA"],
            state="readonly",
            width=40
        )
        tipo_combo.grid(row=current_row, column=1, sticky=W, padx=10, pady=5)
        current_row += 1
        
        # Data - Usar Entry simples com formato em vez de DateEntry
        ttk.Label(form_frame, text="Data:").grid(row=current_row, column=0, sticky=W, padx=10, pady=5)
        data_var = ttk.StringVar(value=datetime.now().strftime("%d/%m/%Y"))
        data_entry = ttk.Entry(
            form_frame,
            textvariable=data_var,
            width=20
        )
        data_entry.grid(row=current_row, column=1, sticky=W, padx=10, pady=5)
        current_row += 1
        
        # Cliente
        ttk.Label(form_frame, text="Nome do Cliente:").grid(row=current_row, column=0, sticky=W, padx=10, pady=5)
        cliente_entry = ttk.Entry(form_frame, width=40)
        cliente_entry.grid(row=current_row, column=1, sticky=W, padx=10, pady=5)
        current_row += 1
        
        # Máquina
        ttk.Label(form_frame, text="Tipo de Máquina:").grid(row=current_row, column=0, sticky=W, padx=10, pady=5)
        maquina_entry = ttk.Entry(form_frame, width=40)
        maquina_entry.grid(row=current_row, column=1, sticky=W, padx=10, pady=5)
        current_row += 1
        
        # Patrimônio
        ttk.Label(form_frame, text="Nº de Patrimônio:").grid(row=current_row, column=0, sticky=W, padx=10, pady=5)
        patrimonio_entry = ttk.Entry(form_frame, width=40)
        patrimonio_entry.grid(row=current_row, column=1, sticky=W, padx=10, pady=5)
        current_row += 1
        
        # Número de Série
        ttk.Label(form_frame, text="Nº de Série:").grid(row=current_row, column=0, sticky=W, padx=10, pady=5)
        numero_serie_entry = ttk.Entry(form_frame, width=40)
        numero_serie_entry.grid(row=current_row, column=1, sticky=W, padx=10, pady=5)
        current_row += 1
        
        # SEÇÃO 2: LOCALIZAÇÃO
        ttk.Label(form_frame, text="LOCALIZAÇÃO", font=("Helvetica", 12, "bold")).grid(row=current_row, column=0, columnspan=2, sticky=W, padx=10, pady=(20,5))
        current_row += 1
        ttk.Separator(form_frame, orient=HORIZONTAL).grid(row=current_row, column=0, columnspan=2, sticky=EW, padx=10, pady=5)
        current_row += 1
        
        # Local
        ttk.Label(form_frame, text="Local:").grid(row=current_row, column=0, sticky=W, padx=10, pady=5)
        local_entry = ttk.Entry(form_frame, width=40)
        local_entry.grid(row=current_row, column=1, sticky=W, padx=10, pady=5)
        current_row += 1
        
        # Endereço
        ttk.Label(form_frame, text="Endereço:").grid(row=current_row, column=0, sticky=W, padx=10, pady=5)
        endereco_entry = ttk.Entry(form_frame, width=40)
        endereco_entry.grid(row=current_row, column=1, sticky=W, padx=10, pady=5)
        current_row += 1
        
        # Cidade (NOVO CAMPO)
        ttk.Label(form_frame, text="Cidade:").grid(row=current_row, column=0, sticky=W, padx=10, pady=5)
        cidade_entry = ttk.Entry(form_frame, width=40)
        cidade_entry.grid(row=current_row, column=1, sticky=W, padx=10, pady=5)
        current_row += 1
        
        # Telefone
        ttk.Label(form_frame, text="Telefone:").grid(row=current_row, column=0, sticky=W, padx=10, pady=5)
        telefone_entry = ttk.Entry(form_frame, width=40)
        telefone_entry.grid(row=current_row, column=1, sticky=W, padx=10, pady=5)
        current_row += 1
        
        # CEP
        ttk.Label(form_frame, text="CEP:").grid(row=current_row, column=0, sticky=W, padx=10, pady=5)
        cep_entry = ttk.Entry(form_frame, width=40)
        cep_entry.grid(row=current_row, column=1, sticky=W, padx=10, pady=5)
        current_row += 1
        
        # SEÇÃO 3: CONTATO
        ttk.Label(form_frame, text="INFORMAÇÕES DE CONTATO", font=("Helvetica", 12, "bold")).grid(row=current_row, column=0, columnspan=2, sticky=W, padx=10, pady=(20,5))
        current_row += 1
        ttk.Separator(form_frame, orient=HORIZONTAL).grid(row=current_row, column=0, columnspan=2, sticky=EW, padx=10, pady=5)
        current_row += 1
        
        # Nome do Contato
        ttk.Label(form_frame, text="Nome do Contato:").grid(row=current_row, column=0, sticky=W, padx=10, pady=5)
        contato_nome_entry = ttk.Entry(form_frame, width=40)
        contato_nome_entry.grid(row=current_row, column=1, sticky=W, padx=10, pady=5)
        current_row += 1
        
        # Telefone do Contato 1
        ttk.Label(form_frame, text="Telefone do Contato 1:").grid(row=current_row, column=0, sticky=W, padx=10, pady=5)
        contato_telefone1_entry = ttk.Entry(form_frame, width=40)
        contato_telefone1_entry.grid(row=current_row, column=1, sticky=W, padx=10, pady=5)
        current_row += 1
        
        # Telefone do Contato 2
        ttk.Label(form_frame, text="Telefone do Contato 2:").grid(row=current_row, column=0, sticky=W, padx=10, pady=5)
        contato_telefone2_entry = ttk.Entry(form_frame, width=40)
        contato_telefone2_entry.grid(row=current_row, column=1, sticky=W, padx=10, pady=5)
        current_row += 1
        
        # SEÇÃO 4: LOGÍSTICA
        ttk.Label(form_frame, text="LOGÍSTICA E PRAZOS", font=("Helvetica", 12, "bold")).grid(row=current_row, column=0, columnspan=2, sticky=W, padx=10, pady=(20,5))
        current_row += 1
        ttk.Separator(form_frame, orient=HORIZONTAL).grid(row=current_row, column=0, columnspan=2, sticky=EW, padx=10, pady=5)
        current_row += 1
        
        # Prazo - Remover DateEntry e usar apenas o Entry
        ttk.Label(form_frame, text="Prazo de Entrega:").grid(row=current_row, column=0, sticky=W, padx=10, pady=5)
        prazo_var = ttk.StringVar(value=(datetime.now() + timedelta(days=7)).strftime("%d/%m/%Y"))
        prazo_entry = ttk.Entry(
            form_frame,
            textvariable=prazo_var,
            width=20
        )
        prazo_entry.grid(row=current_row, column=1, sticky=W, padx=10, pady=5)
        current_row += 1
        
        # Viagem
        ttk.Label(form_frame, text="Necessita de Viagem?").grid(row=current_row, column=0, sticky=W, padx=10, pady=5)
        viagem_var = ttk.StringVar(value="Não")
        viagem_frame = ttk.Frame(form_frame)
        viagem_frame.grid(row=current_row, column=1, sticky=W, padx=10, pady=5)
        ttk.Radiobutton(viagem_frame, text="Sim", value="Sim", variable=viagem_var).pack(side=LEFT, padx=5)
        ttk.Radiobutton(viagem_frame, text="Não", value="Não", variable=viagem_var).pack(side=LEFT, padx=5)
        current_row += 1
        
        # Hospedagem
        ttk.Label(form_frame, text="Tipo de Hospedagem:").grid(row=current_row, column=0, sticky=W, padx=10, pady=5)
        hospedagem_var = ttk.StringVar(value="Não necessário")
        hospedagem_combo = ttk.Combobox(
            form_frame,
            textvariable=hospedagem_var,
            values=["Não necessário", "Hotel", "Airbnb", "Outro"],
            state="readonly",
            width=40
        )
        hospedagem_combo.grid(row=current_row, column=1, sticky=W, padx=10, pady=5)
        current_row += 1
        
        # SEÇÃO 5: DESCRIÇÕES
        ttk.Label(form_frame, text="DESCRIÇÕES", font=("Helvetica", 12, "bold")).grid(row=current_row, column=0, columnspan=2, sticky=W, padx=10, pady=(20,5))
        current_row += 1
        ttk.Separator(form_frame, orient=HORIZONTAL).grid(row=current_row, column=0, columnspan=2, sticky=EW, padx=10, pady=5)
        current_row += 1
        
        # Descrição do Problema
        ttk.Label(form_frame, text="Descrição do Problema:").grid(row=current_row, column=0, sticky=W, padx=10, pady=5)
        current_row += 1
        descricao_text = ttk.Text(form_frame, width=60, height=5)
        descricao_text.grid(row=current_row, column=0, columnspan=2, sticky=W, padx=10, pady=5)
        current_row += 1
        
        # Descrição do Serviço
        ttk.Label(form_frame, text="Descrição do Serviço:").grid(row=current_row, column=0, sticky=W, padx=10, pady=5)
        current_row += 1
        descricao_servico_text = ttk.Text(form_frame, width=60, height=5)
        descricao_servico_text.grid(row=current_row, column=0, columnspan=2, sticky=W, padx=10, pady=5)
        current_row += 1
        
        # Urgência
        ttk.Label(form_frame, text="Nível de Urgência:").grid(row=current_row, column=0, sticky=W, padx=10, pady=5)
        urgencia_var = ttk.StringVar(value="Normal")
        urgencia_combo = ttk.Combobox(
            form_frame,
            textvariable=urgencia_var,
            values=["Normal", "Alta", "Crítica"],
            state="readonly",
            width=40
        )
        urgencia_combo.grid(row=current_row, column=1, sticky=W, padx=10, pady=5)
        current_row += 1
        
        # Botões
        button_frame = ttk.Frame(form_frame)
        button_frame.grid(row=current_row, column=0, columnspan=2, pady=20)
        
        def salvar_os():
            # Validar campos obrigatórios
            if not empresa_var.get().strip():
                mostrar_mensagem(nova_janela, "Erro", "Selecionar a empresa é obrigatório!", "erro")
                return
            if not cliente_entry.get().strip():
                mostrar_mensagem(nova_janela, "Erro", "Nome do cliente é obrigatório!", "erro")
                return
            if not maquina_entry.get().strip():
                mostrar_mensagem(nova_janela, "Erro", "Tipo de máquina é obrigatório!", "erro")
                return
            if not descricao_text.get("1.0", END).strip():
                mostrar_mensagem(nova_janela, "Erro", "Descrição do problema é obrigatória!", "erro")
                return
            
            # Gerar número da OS
            ano = datetime.now().strftime("%y")
            conn = sqlite3.connect('moraca.db')
            c = conn.cursor()
            c.execute("SELECT COUNT(*) FROM os")
            count = c.fetchone()[0]
            numero = f"OS{ano}{str(count + 1).zfill(3)}"
            
            # Preparar dados
            cliente = cliente_entry.get()
            maquina = maquina_entry.get()
            descricao = descricao_text.get("1.0", END)
            urgencia = urgencia_var.get()
            tipo = tipo_var.get()
            data = data_var.get()
            status = "ABERTA"  # Status padrão é ABERTA
            
            # Novos campos
            patrimonio = patrimonio_entry.get()
            numero_serie = numero_serie_entry.get()
            local = local_entry.get()
            endereco = endereco_entry.get()
            cidade = cidade_entry.get()  # Obter o valor do campo cidade
            telefone = telefone_entry.get()
            cep = cep_entry.get()
            contato_nome = contato_nome_entry.get()
            contato_telefone1 = contato_telefone1_entry.get()
            contato_telefone2 = contato_telefone2_entry.get()
            descricao_servico = descricao_servico_text.get("1.0", END)
            
            # Campos adicionais
            necessita_viagem = viagem_var.get()
            tipo_hospedagem = hospedagem_var.get()
            prazo_entrega = prazo_var.get()
            empresa = empresa_var.get()
            
            # Inserir no banco
            c.execute('''INSERT INTO os (numero, cliente, maquina, descricao, urgencia, 
                data, status, tipo, patrimonio, numero_serie, local, endereco, cidade, 
                telefone, cep, contato_nome, contato_telefone1, contato_telefone2, 
                descricao_servico, necessita_viagem, tipo_hospedagem, prazo_entrega, empresa) 
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
                     (numero, cliente, maquina, descricao, urgencia, data, status, tipo,
                      patrimonio, numero_serie, local, endereco, cidade, telefone, cep,
                      contato_nome, contato_telefone1, contato_telefone2, descricao_servico,
                      necessita_viagem, tipo_hospedagem, prazo_entrega, empresa))
            conn.commit()
            conn.close()
            
            # Criar dicionário com os dados do cliente para o modelo Ziehm
            dados_cliente = {
                'cliente': cliente,
                'equipamento': maquina,
                'modelo': maquina,
                'numero_serie': numero_serie,
                'patrimonio': patrimonio,
                'endereco': endereco,
                'contato': contato_nome,
                'telefone': telefone,
                'telefone1': contato_telefone1,
                'telefone2': contato_telefone2,
                'local': local,
                'cep': cep,
                'descricao': descricao,
                'descricao_servico': descricao_servico,
                'urgencia': urgencia,
                'tipo': tipo,
                'data': data,
                'prazo_entrega': prazo_entrega,
                'necessita_viagem': necessita_viagem,
                'tipo_hospedagem': tipo_hospedagem,
                'status': status
            }
            
            # Gerar arquivos Excel - Verificar qual modelo usar
            if empresa == "ZIEHM":
                # Usar o modelo da Ziehm
                criar_os_ziehm(numero, dados_cliente)
            else:
                # Usar o modelo padrão da Moraca
                self.gerar_arquivo_excel_os(numero, cliente, maquina, descricao, urgencia, data, status, tipo,
                                    patrimonio, numero_serie, local, endereco, cidade, telefone, cep,
                                    contato_nome, contato_telefone1, contato_telefone2, descricao_servico,
                                    necessita_viagem, tipo_hospedagem, prazo_entrega, empresa)

            # Gerar arquivos Word (sempre usa o mesmo modelo)
            self.gerar_arquivo_os(numero, cliente, maquina, descricao, urgencia, data, status, tipo,
                                patrimonio, numero_serie, local, endereco, cidade, telefone, cep,
                                contato_nome, contato_telefone1, contato_telefone2, descricao_servico,
                                necessita_viagem, tipo_hospedagem, prazo_entrega, empresa)
            
            # Mostrar mensagem de sucesso
            mostrar_mensagem(nova_janela, "Sucesso", f"OS {numero} criada com sucesso!", "sucesso")
            
            self.atualizar_lista()
            nova_janela.destroy()
        
        # Botões
        ttk.Button(
            button_frame,
            text="Criar OS",
            command=salvar_os,
            style="primary.TButton",
            width=15
        ).pack(side=LEFT, padx=5)
    
        ttk.Button(
            button_frame,
            text="Cancelar",
            command=nova_janela.destroy,
            style="secondary.TButton",
            width=15
        ).pack(side=LEFT, padx=5)
        
        # Atualizar a região de scroll no final
        form_frame.update_idletasks()
        canvas.configure(scrollregion=canvas.bbox("all"))
        
        # Garantir que a rolagem funciona
        canvas.bind_all("<MouseWheel>", lambda event: canvas.yview_scroll(int(-1*(event.delta/120)), "units"))
    
    def buscar_os(self):
        termo = self.search_entry.get().strip()
        # Limpar tabela
        for item in self.tabela.get_children():
            self.tabela.delete(item)
        
        # Se o termo estiver vazio, mostrar as últimas OS
        if not termo:
            self.atualizar_lista()
            return
        
        # Buscar no banco com limite maior
        conn = sqlite3.connect('moraca.db')
        c = conn.cursor()
        c.execute('''SELECT numero, cliente, data, status FROM os
                    WHERE numero LIKE ? OR cliente LIKE ?
                    ORDER BY data DESC LIMIT 30''', (f"%{termo}%", f"%{termo}%"))
        
        registros = c.fetchall()
        conn.close()
        
        # Verificar se encontrou resultados
        if not registros:
            # Não há resultados, mostrar mensagem
            mostrar_mensagem(self.root, "Busca", f"Nenhuma OS encontrada para '{termo}'", "aviso")
            # Recarregar a lista original
            self.atualizar_lista()
            return
        
        # Inserir os resultados na tabela
        for row in registros:
            self.tabela.insert("", END, values=row, tags=("encontrado",))
        
        # Estilizar os resultados encontrados
        self.tabela.tag_configure("encontrado", background=self.cores["encontrado"])
        
        # Mostrar mensagem com o número de resultados
        mostrar_mensagem(self.root, "Busca", f"Foram encontradas {len(registros)} OS para '{termo}'", "sucesso")
    
    def atualizar_lista(self):
        # Limpar tabela
        for item in self.tabela.get_children():
            self.tabela.delete(item)
        
        # Buscar últimas OS (aumentando o limite para mostrar mais)
        conn = sqlite3.connect('moraca.db')
        c = conn.cursor()
        c.execute('''SELECT numero, cliente, data, status FROM os
                    ORDER BY data DESC LIMIT 20''')
        
        for row in c.fetchall():
            # Converter status antigos para os novos formatos se necessário
            numero, cliente, data, status = row
            # Verificar se o status está no formato antigo e converter
            if status in self.mapeamento_status:
                status = self.mapeamento_status[status]
            
            # Adicionar tags baseadas no status para colorir as linhas
            if status == "ABERTA":
                self.tabela.insert("", END, values=(numero, cliente, data, status), tags=("aberta",))
            elif status == "FECHADA":
                self.tabela.insert("", END, values=(numero, cliente, data, status), tags=("fechada",))
            else:
                # Se por algum motivo o status não estiver nos formatos esperados
                self.tabela.insert("", END, values=(numero, cliente, data, status))
        
        conn.close()
        
        # Configurar cores para os diferentes status
        self.tabela.tag_configure("aberta", background=self.cores["aberta"])
        self.tabela.tag_configure("fechada", background=self.cores["fechada"])
    
    def mostrar_todas_os(self):
        """Exibe todas as OS no sistema."""
        # Limpar tabela
        for item in self.tabela.get_children():
            self.tabela.delete(item)
        
        # Buscar todas as OS
        conn = sqlite3.connect('moraca.db')
        c = conn.cursor()
        c.execute('''SELECT numero, cliente, data, status FROM os
                    ORDER BY data DESC''')
        
        registros = c.fetchall()
        conn.close()
        
        # Verificar se existem OS
        if not registros:
            mostrar_mensagem(self.root, "Informação", "Não existem OS cadastradas no sistema.", "info")
            return
        
        # Inserir todas as OS na tabela
        for row in registros:
            # Converter status antigos para os novos formatos se necessário
            numero, cliente, data, status = row
            # Verificar se o status está no formato antigo e converter
            if status in self.mapeamento_status:
                status = self.mapeamento_status[status]
            
            # Adicionar tags baseadas no status para colorir as linhas
            if status == "ABERTA":
                self.tabela.insert("", END, values=(numero, cliente, data, status), tags=("aberta",))
            elif status == "FECHADA":
                self.tabela.insert("", END, values=(numero, cliente, data, status), tags=("fechada",))
            else:
                # Se por algum motivo o status não estiver nos formatos esperados
                self.tabela.insert("", END, values=(numero, cliente, data, status))
        
        # Configurar cores para os diferentes status
        self.tabela.tag_configure("aberta", background=self.cores["aberta"])
        self.tabela.tag_configure("fechada", background=self.cores["fechada"])
        
        # Mostrar mensagem com o número total de OS
        mostrar_mensagem(self.root, "Lista Completa", f"Exibindo todas as {len(registros)} OS do sistema.", "info")
    
    def adicionar_cabecalho(self, doc, numero_os):
        """Adiciona o cabeçalho com o timbre da empresa no documento"""
        # Adicionar tabela para o cabeçalho
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.allow_autofit = False
        
        # Células da tabela
        cell_os = table.cell(0, 0)
        cell_logo = table.cell(0, 1)
        
        # Configurar larguras
        cell_os.width = Inches(2.5)
        cell_logo.width = Inches(4.0)
        
        # Adicionar número da OS
        os_paragraph = cell_os.paragraphs[0]
        os_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        os_run = os_paragraph.add_run()
        os_run.add_text("Ordem de Serviço\n")
        os_run.font.size = Pt(12)
        os_run.font.bold = True
        os_run.add_text(f"{numero_os}")
        os_run.font.size = Pt(14)
        os_run.font.bold = True
        
        # Adicionar logo
        logo_paragraph = cell_logo.paragraphs[0]
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        logo_run = logo_paragraph.add_run()
        logo_path = os.path.join("assets", "logo.png")
        if os.path.exists(logo_path):
            logo_run.add_picture(logo_path, width=Inches(3.5))
        
        # Adicionar espaço após o cabeçalho
        doc.add_paragraph()

    def gerar_os_interna_andamento(self):
        # Criar uma janela para solicitar o número da OS
        janela = ttk.Toplevel(self.root)
        janela.title("Gerar OS Interna de Andamento")
        janela.geometry("400x200")
        
        # Frame principal
        frame = ttk.Frame(janela, padding=20)
        frame.pack(fill=BOTH, expand=YES)
        
        # Título
        ttk.Label(
            frame,
            text="Gerar OS Interna de Andamento",
            font=("Helvetica", 12, "bold")
        ).pack(pady=(0, 20))
        
        # Frame para entrada
        input_frame = ttk.Frame(frame)
        input_frame.pack(pady=10)
        
        # Campo para número da OS
        ttk.Label(input_frame, text="Número da OS:").grid(row=0, column=0, padx=5, pady=5, sticky=W)
        numero_os_entry = ttk.Entry(input_frame, width=20)
        numero_os_entry.grid(row=0, column=1, padx=5, pady=5, sticky=W)
        
        # Adicionar instruções sobre o formato
        ttk.Label(input_frame, text="Formato: ano + número (ex: 25019)", foreground="gray").grid(row=1, column=0, columnspan=2, padx=5, pady=0, sticky=W)
        
        # Função para gerar o documento
        def gerar_documento():
            numero_os = numero_os_entry.get().strip()
            
            if not numero_os:
                mostrar_mensagem(janela, "Erro", "Por favor, informe o número da OS.", "erro")
                return
            
            try:
                # Caminho para salvar o documento
                pasta_destino = os.path.join("/home/gabriel-balsarin/Projetos/Moraca/MORACA", "MORACA1", "DOCUMENTOS", "tecnico", "andamento")
                
                # Verificar se o módulo está disponível
                if 'criar_os_interna_andamento' in globals():
                    # Gerar o documento
                    caminho_arquivo = criar_os_interna_andamento(numero_os, pasta_destino)
                    
                    mostrar_mensagem(janela, "Sucesso", f"Documento gerado com sucesso em:\n{caminho_arquivo}", "sucesso")
                    janela.destroy()
                else:
                    mostrar_mensagem(janela, "Erro", "Funcionalidade não disponível. O módulo os_interna_andamento não está acessível.", "erro")
            except Exception as e:
                mostrar_mensagem(janela, "Erro", f"Ocorreu um erro ao gerar o documento:\n{str(e)}", "erro")
        
        # Frame para botões
        button_frame = ttk.Frame(frame)
        button_frame.pack(pady=20)
        
        # Botões
        ttk.Button(
            button_frame,
            text="Gerar",
            command=gerar_documento,
            style="primary.TButton",
            width=15
        ).pack(side=LEFT, padx=5)
        
        ttk.Button(
            button_frame,
            text="Cancelar",
            command=janela.destroy,
            style="secondary.TButton",
            width=15
        ).pack(side=LEFT, padx=5)

    def abrir_aba_impressao(self, event):
        """Abre uma nova aba com opções de impressão para a OS selecionada"""
        # Obter o item selecionado
        item = self.tabela.selection()
        if not item:
            mostrar_mensagem(self.root, "Aviso", "Selecione uma OS para imprimir", "aviso")
            return
        
        # Obter os dados da OS selecionada
        os_data = self.tabela.item(item[0], 'values')
        numero_os = os_data[0]  # Primeira coluna (número)
        cliente = os_data[1]    # Segunda coluna (cliente)
        data = os_data[2]       # Terceira coluna (data)
        status_atual = os_data[3]  # Quarta coluna (status)
        
        # Converter status antigo para novo formato, se necessário
        if status_atual in self.mapeamento_status:
            status_atual = self.mapeamento_status[status_atual]
        
        # Buscar dados completos do banco
        conn = sqlite3.connect('moraca.db')
        c = conn.cursor()
        c.execute("SELECT * FROM os WHERE numero = ?", (numero_os,))
        os_completa = c.fetchone()
        conn.close()
        
        if not os_completa:
            mostrar_mensagem(self.root, "Erro", f"Não foi possível encontrar os dados da OS {numero_os}", "erro")
            return
        
        # Criar nova janela para a aba de impressão
        janela_impressao = ttk.Toplevel(self.root)
        janela_impressao.title(f"Opções de Impressão - OS {numero_os}")
        janela_impressao.geometry("700x600")
        
        # Frame principal
        frame_principal = ttk.Frame(janela_impressao, padding=20)
        frame_principal.pack(fill=BOTH, expand=YES)
        
        # Título
        ttk.Label(
            frame_principal,
            text=f"Gerenciar OS {numero_os}",
            font=("Helvetica", 14, "bold")
        ).pack(pady=(0, 20))
        
        # Informações da OS
        info_frame = ttk.Frame(frame_principal)
        info_frame.pack(fill=X, pady=10)
        
        ttk.Label(
            info_frame,
            text=f"Cliente: {cliente}",
            font=("Helvetica", 12)
        ).pack(anchor="w")
        
        ttk.Label(
            info_frame,
            text=f"Data: {data}",
            font=("Helvetica", 12)
        ).pack(anchor="w")
        
        # Frame para gestão de status
        status_frame = ttk.LabelFrame(frame_principal, text="Status da OS", padding=10)
        status_frame.pack(fill=X, pady=10)
        
        # Mostrar status atual
        ttk.Label(
            status_frame,
            text=f"Status Atual: {status_atual}",
            font=("Helvetica", 11, "bold")
        ).pack(anchor="w", pady=5)
        
        # Opções de status
        status_options = ["ABERTA", "FECHADA"]
        status_var = ttk.StringVar(value=status_atual)
        
        # Frame para o combobox e botão
        status_selection_frame = ttk.Frame(status_frame)
        status_selection_frame.pack(fill=X, pady=5)
        
        ttk.Label(
            status_selection_frame,
            text="Alterar para:",
            font=("Helvetica", 10)
        ).pack(side=LEFT, padx=(0, 10))
        
        status_combobox = ttk.Combobox(
            status_selection_frame,
            textvariable=status_var,
            values=status_options,
            state="readonly",
            width=20
        )
        status_combobox.pack(side=LEFT, padx=5)
        
        # Função para alterar o status
        def alterar_status():
            novo_status = status_var.get()
            if novo_status == status_atual:
                mostrar_mensagem(janela_impressao, "Aviso", "O status selecionado é o mesmo que o atual.", "aviso")
                return
            
            # Perguntar ao usuário para confirmar
            confirmar = mostrar_mensagem(
                janela_impressao, 
                "Confirmar Alteração", 
                f"Deseja realmente alterar o status da OS {numero_os} de '{status_atual}' para '{novo_status}'?",
                "aviso"
            )
            
            if not confirmar:
                return
            
            try:
                # Atualizar o banco de dados
                conn = sqlite3.connect('moraca.db')
                c = conn.cursor()
                c.execute("UPDATE os SET status = ? WHERE numero = ?", (novo_status, numero_os))
                conn.commit()
                
                # Obter o cliente para localizarmos a pasta da OS
                c.execute("SELECT cliente FROM os WHERE numero = ?", (numero_os,))
                resultado = c.fetchone()
                conn.close()
                
                if resultado and resultado[0]:
                    cliente = resultado[0]
                    
                    # Mover a pasta da OS para "OS fechada" quando o status for FECHADA
                    if novo_status == "FECHADA":
                        # Obter o caminho atual da pasta da OS
                        pasta_origem = self.get_pasta_os_cliente(numero_os, cliente)
                        
                        # Criar caminho para a pasta "OS fechada" mantendo a mesma estrutura
                        base_path = os.path.join("MORACA", "MORACA1", "DOCUMENTOS", "tecnico", "os")
                        
                        # Extrair ano da OS
                        match = re.search(r'OS(\d{2})(\d+)', numero_os)
                        if match:
                            ano = match.group(1)
                            ano_completo = f"20{ano}"
                        else:
                            ano_completo = datetime.now().strftime("%Y")
                            
                        # Criar estrutura na pasta "OS fechada"
                        pasta_os_fechada = os.path.join(base_path, "OS-fechada", f"OS-{ano_completo}")
                        
                        # Garantir que o diretório de destino exista
                        if not os.path.exists(pasta_os_fechada):
                            os.makedirs(pasta_os_fechada)
                            
                        # Extrair o nome da pasta da OS do caminho origem
                        nome_pasta_os = os.path.basename(pasta_origem)
                        
                        # Caminho de destino completo
                        pasta_destino = os.path.join(pasta_os_fechada, nome_pasta_os)
                        
                        # Verificar se a pasta origem existe
                        if os.path.exists(pasta_origem):
                            # Se o destino já existir, precisamos lidar com isso
                            if os.path.exists(pasta_destino):
                                # Adicionar timestamp para evitar duplicidade
                                timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
                                pasta_destino = os.path.join(pasta_os_fechada, f"{nome_pasta_os}_{timestamp}")
                            
                            try:
                                # Copiar a pasta inteira com seu conteúdo
                                shutil.copytree(pasta_origem, pasta_destino)
                                
                                # Remover pasta original após copiar
                                shutil.rmtree(pasta_origem)
                                
                                mostrar_mensagem(
                                    janela_impressao, 
                                    "Sucesso", 
                                    f"Pasta da OS {numero_os} movida para 'OS fechada'.",
                                    "sucesso"
                                )
                            except Exception as e:
                                mostrar_mensagem(
                                    janela_impressao, 
                                    "Aviso", 
                                    f"Status alterado, mas não foi possível mover a pasta: {str(e)}",
                                    "aviso"
                                )
                        
                    # Código existente para mover entre pendentes e concluídas, se ainda necessário
                    # [...]
                
                # Atualizar a lista na tela principal
                self.atualizar_lista()
                
                # Mostrar mensagem de sucesso
                mostrar_mensagem(
                    janela_impressao, 
                    "Sucesso", 
                    f"Status da OS {numero_os} alterado para '{novo_status}'.",
                    "sucesso"
                )
                
                # Fechar a janela após a alteração
                janela_impressao.destroy()
                
            except Exception as e:
                mostrar_mensagem(
                    janela_impressao, 
                    "Erro", 
                    f"Erro ao alterar status: {str(e)}",
                    "erro"
                )
                
        # Botão para alterar status
        ttk.Button(
            status_selection_frame,
            text="Alterar Status",
            command=alterar_status,
            style="primary.TButton",
            width=15
        ).pack(side=LEFT, padx=10)
        
        # Separador
        ttk.Separator(frame_principal).pack(fill=X, pady=20)
        
        # Frame para botões de impressão
        buttons_frame = ttk.Frame(frame_principal)
        buttons_frame.pack(pady=10)
        
        ttk.Label(
            buttons_frame,
            text="Opções de Impressão",
            font=("Helvetica", 12, "bold")
        ).pack(pady=10)
        
        # Estilo para botões
        button_style = "primary.TButton"
        button_width = 25
        
        # Botão para imprimir OS de visita
        ttk.Button(
            buttons_frame,
            text="Imprimir OS de Visita",
            command=lambda: [self.imprimir_os_visita(numero_os), janela_impressao.destroy()],
            style=button_style,
            width=button_width
        ).pack(pady=5)
        
        # Botão para imprimir OS Andamento
        ttk.Button(
            buttons_frame,
            text="Imprimir OS andamento",
            command=lambda: [self.imprimir_os_andamento(numero_os), janela_impressao.destroy()],
            style=button_style,
            width=button_width
        ).pack(pady=5)
        
        # Botão para imprimir OS preventiva
        ttk.Button(
            buttons_frame,
            text="Imprimir OS de Preventiva",
            command=lambda: [self.imprimir_os_preventiva(numero_os), janela_impressao.destroy()],
            style=button_style,
            width=button_width
        ).pack(pady=5)
        
        # Botão para imprimir OS interna
        ttk.Button(
            buttons_frame,
            text="Imprimir OS Interna",
            command=lambda: [self.imprimir_os_interna(numero_os), janela_impressao.destroy()],
            style=button_style,
            width=button_width
        ).pack(pady=5)
        
        # Botão para imprimir OS Ziehm
        ttk.Button(
            buttons_frame,
            text="Imprimir OS Ziehm",
            command=lambda: [self.imprimir_os_ziehm(numero_os), janela_impressao.destroy()],
            style=button_style,
            width=button_width
        ).pack(pady=5)
        
        # Espaçador
        ttk.Label(frame_principal, text="").pack(pady=5)
        
        # Botão de fechar
        ttk.Button(
            frame_principal,
            text="Fechar",
            command=janela_impressao.destroy,
            style="secondary.TButton",
            width=15
        ).pack(pady=10)

    def imprimir_os_visita(self, numero_os):
        """Gera e abre documento de OS de visita"""
        try:
            # Verificar se o arquivo já existe
            caminho_arquivo = self.get_caminho_arquivo('visita', numero_os)
            
            if caminho_arquivo and os.path.exists(caminho_arquivo):
                # Se o arquivo já existe, apenas abre sem recriar
                self.abrir_arquivo(caminho_arquivo)
                mostrar_mensagem(self.root, "Sucesso", f"Documento existente aberto em:\n{caminho_arquivo}", "sucesso")
                return
                
            # Se não existe, tentar criar - por enquanto apenas mostra mensagem de desenvolvimento
            mostrar_mensagem(self.root, "Em Desenvolvimento", "Impressão de OS de Visita em desenvolvimento.")
        except Exception as e:
            mostrar_mensagem(self.root, "Erro", f"Erro ao gerar OS de Visita: {str(e)}", "erro")
    
    def imprimir_os_andamento(self, numero_os):
        """Gera e abre documento de OS de andamento"""
        try:
            # Buscar os dados da OS primeiro para ter informações do cliente
            conn = sqlite3.connect('moraca.db')
            c = conn.cursor()
            c.execute("SELECT cliente FROM os WHERE numero = ?", (numero_os,))
            resultado = c.fetchone()
            conn.close()
            
            if not resultado:
                mostrar_mensagem(self.root, "Erro", f"Não foi possível encontrar os dados da OS {numero_os}", "erro")
                return
                
            cliente = resultado[0]
            
            # Verificar primeiro na nova estrutura de pastas
            pasta_os = self.get_pasta_os_cliente(numero_os, cliente)
            arquivo_excel = os.path.join(pasta_os, f"{numero_os}.xlsx")
            
            # Verificar se o arquivo existe na nova estrutura
            if os.path.exists(arquivo_excel):
                # Se o arquivo já existe, apenas abre sem recriar
                self.abrir_arquivo(arquivo_excel)
                mostrar_mensagem(self.root, "Sucesso", f"Documento existente aberto em:\n{arquivo_excel}", "sucesso")
                return
                
            # Verificar se o arquivo já existe no caminho antigo
            caminho_arquivo = self.get_caminho_arquivo('andamento', numero_os)
            if caminho_arquivo and os.path.exists(caminho_arquivo):
                # Se o arquivo já existe no local antigo, copiar para a nova estrutura
                os.makedirs(os.path.dirname(arquivo_excel), exist_ok=True)
                shutil.copy2(caminho_arquivo, arquivo_excel)
                
                # Abrir o arquivo da nova localização
                self.abrir_arquivo(arquivo_excel)
                mostrar_mensagem(self.root, "Sucesso", f"Documento existente copiado e aberto em:\n{arquivo_excel}", "sucesso")
                return
                
            # Se não existe, continua com a criação do documento
            # Gerar o documento usando a função existente
            caminho_arquivo = criar_os_interna_andamento(numero_os)
            
            # Copiar para a nova estrutura
            os.makedirs(os.path.dirname(arquivo_excel), exist_ok=True)
            shutil.copy2(caminho_arquivo, arquivo_excel)
            
            # Abrir o documento com o aplicativo padrão (já da nova localização)
            self.abrir_arquivo(arquivo_excel)
            
            mostrar_mensagem(self.root, "Sucesso", f"OS de Andamento gerada com sucesso em:\n{arquivo_excel}", "sucesso")
        except Exception as e:
            mostrar_mensagem(self.root, "Erro", f"Erro ao gerar OS de Andamento: {str(e)}", "erro")
    
    def imprimir_os_preventiva(self, numero_os):
        """Gera e abre documento de OS preventiva"""
        try:
            # Verificar se o arquivo já existe
            caminho_arquivo = self.get_caminho_arquivo('preventiva', numero_os)
            
            if caminho_arquivo and os.path.exists(caminho_arquivo):
                # Se o arquivo já existe, apenas abre sem recriar
                self.abrir_arquivo(caminho_arquivo)
                mostrar_mensagem(self.root, "Sucesso", f"Documento existente aberto em:\n{caminho_arquivo}", "sucesso")
                return
                
            # Se não existe, tentar criar - por enquanto apenas mostra mensagem de desenvolvimento
            mostrar_mensagem(self.root, "Em Desenvolvimento", "Impressão de OS Preventiva em desenvolvimento.")
        except Exception as e:
            mostrar_mensagem(self.root, "Erro", f"Erro ao gerar OS Preventiva: {str(e)}", "erro")
    
    def imprimir_os_interna(self, numero_os):
        """Gera e abre documento de OS interna"""
        try:
            # Buscar os dados da OS primeiro para ter informações do cliente
            conn = sqlite3.connect('moraca.db')
            c = conn.cursor()
            c.execute("SELECT cliente FROM os WHERE numero = ?", (numero_os,))
            resultado = c.fetchone()
            conn.close()
            
            if not resultado:
                mostrar_mensagem(self.root, "Erro", f"Não foi possível encontrar os dados da OS {numero_os}", "erro")
                return
                
            cliente = resultado[0]
            
            # Verificar primeiro na nova estrutura de pastas
            pasta_os = self.get_pasta_os_cliente(numero_os, cliente)
            arquivo_excel = os.path.join(pasta_os, f"{numero_os}_interna.xlsx")
            
            # Verificar se o arquivo existe na nova estrutura
            if os.path.exists(arquivo_excel):
                # Se o arquivo já existe, apenas abre sem recriar
                self.abrir_arquivo(arquivo_excel)
                mostrar_mensagem(self.root, "Sucesso", f"Documento existente aberto em:\n{arquivo_excel}", "sucesso")
                return
                
            # Verificar se o arquivo já existe no caminho antigo
            caminho_arquivo = self.get_caminho_arquivo('interna', numero_os)
            if caminho_arquivo and os.path.exists(caminho_arquivo):
                # Se o arquivo já existe no local antigo, copiar para a nova estrutura
                os.makedirs(os.path.dirname(arquivo_excel), exist_ok=True)
                shutil.copy2(caminho_arquivo, arquivo_excel)
                
                # Abrir o arquivo da nova localização
                self.abrir_arquivo(arquivo_excel)
                mostrar_mensagem(self.root, "Sucesso", f"Documento existente copiado e aberto em:\n{arquivo_excel}", "sucesso")
                return
                
            # Se não existe, usamos a nova função específica para OS Interna
            # O import já foi feito no início do arquivo
            caminho_arquivo = criar_os_interna(numero_os)
            
            if not caminho_arquivo:
                mostrar_mensagem(self.root, "Erro", f"Não foi possível gerar a OS Interna para {numero_os}", "erro")
                return
            
            # Copiar para a nova estrutura
            os.makedirs(os.path.dirname(arquivo_excel), exist_ok=True)
            shutil.copy2(caminho_arquivo, arquivo_excel)
            
            # Abrir o documento com o aplicativo padrão (já da nova localização)
            self.abrir_arquivo(arquivo_excel)
            
            mostrar_mensagem(self.root, "Sucesso", f"OS Interna gerada com sucesso em:\n{arquivo_excel}", "sucesso")
        except Exception as e:
            mostrar_mensagem(self.root, "Erro", f"Erro ao gerar OS Interna: {str(e)}", "erro")
    
    def imprimir_os_ziehm(self, numero_os):
        """Gera e abre documento de OS Ziehm"""
        try:
            # Buscar dados do cliente para a OS
            conn = sqlite3.connect('moraca.db')
            c = conn.cursor()
            c.execute("SELECT * FROM os WHERE numero = ?", (numero_os,))
            os_data = c.fetchone()
            conn.close()
            
            if not os_data:
                mostrar_mensagem(self.root, "Erro", f"Não foi possível encontrar os dados da OS {numero_os}", "erro")
                return
            
            # Converter dados do BD para dicionário
            colunas = ["id", "numero", "cliente", "maquina", "descricao", "urgencia", 
                      "data", "status", "tipo", "patrimonio", "numero_serie", 
                      "local", "endereco", "telefone", "cep", "contato_nome", 
                      "contato_telefone1", "contato_telefone2", "descricao_servico", 
                      "necessita_viagem", "tipo_hospedagem", "prazo_entrega", "empresa"]
            
            dados_cliente = {colunas[i]: os_data[i] for i in range(len(colunas)) if i < len(os_data)}
            cliente = dados_cliente.get("cliente", "")
            
            # Verificar primeiro na nova estrutura de pastas
            pasta_os = self.get_pasta_os_cliente(numero_os, cliente)
            arquivo_ziehm = os.path.join(pasta_os, f"{numero_os}_ziehm.xlsx")
            
            # Verificar se o arquivo existe na nova estrutura
            if os.path.exists(arquivo_ziehm):
                # Se o arquivo já existe, apenas abre sem recriar
                self.abrir_arquivo(arquivo_ziehm)
                mostrar_mensagem(self.root, "Sucesso", f"Documento Ziehm existente aberto em:\n{arquivo_ziehm}", "sucesso")
                return
                
            # Verificar se o arquivo já existe no caminho antigo
            caminho_arquivo = self.get_caminho_arquivo('ziehm', numero_os)
            if caminho_arquivo and os.path.exists(caminho_arquivo):
                # Se o arquivo já existe no local antigo, copiar para a nova estrutura
                os.makedirs(os.path.dirname(arquivo_ziehm), exist_ok=True)
                shutil.copy2(caminho_arquivo, arquivo_ziehm)
                
                # Abrir o arquivo da nova localização
                self.abrir_arquivo(arquivo_ziehm)
                mostrar_mensagem(self.root, "Sucesso", f"Documento Ziehm existente copiado e aberto em:\n{arquivo_ziehm}", "sucesso")
                return
                
            # Se não existe, continua com a criação do documento
            # Gerar o documento usando a função existente
            caminho_arquivo = criar_os_ziehm(numero_os, dados_cliente)
            
            # Copiar para a nova estrutura
            os.makedirs(os.path.dirname(arquivo_ziehm), exist_ok=True)
            shutil.copy2(caminho_arquivo, arquivo_ziehm)
            
            # Abrir o documento com o aplicativo padrão (já da nova localização)
            self.abrir_arquivo(arquivo_ziehm)
            
            mostrar_mensagem(self.root, "Sucesso", f"OS Ziehm gerada com sucesso em:\n{arquivo_ziehm}", "sucesso")
        except Exception as e:
            mostrar_mensagem(self.root, "Erro", f"Erro ao gerar OS Ziehm: {str(e)}", "erro")
    
    def abrir_arquivo(self, caminho):
        """Abre um arquivo com o aplicativo padrão do sistema"""
        try:
            if platform.system() == 'Windows':
                os.startfile(caminho)
            elif platform.system() == 'Darwin':  # macOS
                subprocess.call(['open', caminho])
            else:  # linux
                subprocess.call(['xdg-open', caminho])
        except Exception as e:
            mostrar_mensagem(self.root, "Erro", f"Erro ao abrir o arquivo: {str(e)}", "erro")

    def imprimir_os_visita_rapida(self):
        """Verifica se uma OS está selecionada e imprime OS de visita"""
        # Obter o item selecionado
        item = self.tabela.selection()
        if not item:
            mostrar_mensagem(self.root, "Aviso", "Selecione uma OS na lista primeiro", "aviso")
            return
        
        # Obter o número da OS selecionada
        numero_os = self.tabela.item(item[0], 'values')[0]
        
        # Chamar a função de impressão
        self.imprimir_os_visita(numero_os)
    
    def imprimir_os_andamento_rapida(self):
        """Verifica se uma OS está selecionada e imprime OS de andamento"""
        # Obter o item selecionado
        item = self.tabela.selection()
        if not item:
            mostrar_mensagem(self.root, "Aviso", "Selecione uma OS na lista primeiro", "aviso")
            return
        
        # Obter o número da OS selecionada
        numero_os = self.tabela.item(item[0], 'values')[0]
        
        # Chamar a função de impressão
        self.imprimir_os_andamento(numero_os)
    
    def imprimir_os_interna_rapida(self):
        """Verifica se uma OS está selecionada e imprime OS interna"""
        # Obter o item selecionado
        item = self.tabela.selection()
        if not item:
            mostrar_mensagem(self.root, "Aviso", "Selecione uma OS na lista primeiro", "aviso")
            return
        
        # Obter o número da OS selecionada
        numero_os = self.tabela.item(item[0], 'values')[0]
        
        # Chamar a função de impressão
        self.imprimir_os_interna(numero_os)
    
    def imprimir_os_ziehm_rapida(self):
        """Verifica se uma OS está selecionada e imprime OS Ziehm"""
        # Obter o item selecionado
        item = self.tabela.selection()
        if not item:
            mostrar_mensagem(self.root, "Aviso", "Selecione uma OS na lista primeiro", "aviso")
            return
        
        # Obter o número da OS selecionada
        numero_os = self.tabela.item(item[0], 'values')[0]
        
        # Chamar a função de impressão
        self.imprimir_os_ziehm(numero_os)
    
    def imprimir_os_preventiva_rapida(self):
        """Verifica se uma OS está selecionada e imprime OS preventiva"""
        # Obter o item selecionado
        item = self.tabela.selection()
        if not item:
            mostrar_mensagem(self.root, "Aviso", "Selecione uma OS na lista primeiro", "aviso")
            return
        
        # Obter o número da OS selecionada
        numero_os = self.tabela.item(item[0], 'values')[0]
        
        # Chamar a função de impressão
        self.imprimir_os_preventiva(numero_os)
    
    def imprimir_os_visita_corretiva_rapida(self):
        """Verifica se uma OS está selecionada e imprime OS de visita corretiva"""
        # Obter o item selecionado
        item = self.tabela.selection()
        if not item:
            mostrar_mensagem(self.root, "Aviso", "Selecione uma OS na lista primeiro", "aviso")
            return
        
        # Obter o número da OS selecionada
        numero_os = self.tabela.item(item[0], 'values')[0]
        
        # Chamar a função de impressão
        self.imprimir_os_visita_corretiva(numero_os)
    
    def imprimir_os_visita_tecnica_rapida(self):
        """Imprime OS de visita técnica do item selecionado na tabela"""
        try:
            # Pegar a linha selecionada na tabela
            selecao = self.tabela.selection()
            if not selecao:
                mostrar_mensagem(self.root, "Aviso", "Selecione uma OS na tabela para imprimir", "aviso")
                return
                
            # Obter o número da OS da linha selecionada
            valores = self.tabela.item(selecao[0], 'values')
            numero_os = valores[0]  # Assumindo que o número da OS está na primeira coluna
            
            # Chamar a função de impressão
            self.imprimir_os_visita_tecnica(numero_os)
        except Exception as e:
            mostrar_mensagem(self.root, "Erro", f"Erro ao imprimir OS de Visita Técnica: {str(e)}", "erro")
    
    def imprimir_os_visita_corretiva(self, numero_os):
        """Gera e abre documento de OS de visita corretiva"""
        try:
            # Verificar se o arquivo já existe
            caminho_arquivo = self.get_caminho_arquivo('visita_corretiva', numero_os)
            
            if caminho_arquivo and os.path.exists(caminho_arquivo):
                # Se o arquivo já existe, apenas abre sem recriar
                self.abrir_arquivo(caminho_arquivo)
                mostrar_mensagem(self.root, "Sucesso", f"Documento existente aberto em:\n{caminho_arquivo}", "sucesso")
                return
                
            # Se não existe, tentar criar - por enquanto apenas mostra mensagem de desenvolvimento
            mostrar_mensagem(self.root, "Em Desenvolvimento", "Impressão de OS de Visita Corretiva em desenvolvimento.")
        except Exception as e:
            mostrar_mensagem(self.root, "Erro", f"Erro ao gerar OS de Visita Corretiva: {str(e)}", "erro")
    
    def imprimir_os_visita_tecnica(self, numero_os):
        """Gera e abre documento de OS de visita técnica"""
        try:
            # Diagnóstico inicial
            print(f"Iniciando impressão da OS de Visita Técnica para número: {numero_os}")
            
            # Formatação padrão do número de OS (do formato "OS xx yyy" para "OSxxxxx")
            numero_formatado_db = None
            
            # Se o número já está no formato "OSxxxxx" sem espaços
            if numero_os.startswith("OS") and " " not in numero_os:
                numero_formatado_db = numero_os
            else:
                # Extrair números
                numeros = re.findall(r'\d+', numero_os)
                if numeros:
                    numeros_juntos = ''.join(numeros)
                    if len(numeros_juntos) >= 5:
                        numero_formatado_db = f"OS{numeros_juntos}"
            
            print(f"Número formatado para consulta no banco: {numero_formatado_db}")
            
            # Buscar os dados da OS primeiro para ter informações do cliente
            conn = sqlite3.connect('moraca.db')
            c = conn.cursor()
            
            # Tentar diferentes formatos de número de OS
            formatos_para_testar = [
                numero_os,  # formato original
                numero_formatado_db  # formato para o banco
            ]
            
            resultado = None
            for formato in formatos_para_testar:
                if formato:
                    print(f"Tentando consulta com: {formato}")
                    c.execute("SELECT * FROM os WHERE numero = ?", (formato,))
                    resultado = c.fetchone()
                    if resultado:
                        print(f"OS encontrada com formato: {formato}")
                        break
            
            if not resultado:
                print(f"OS não encontrada no banco. Verificando números disponíveis...")
                c.execute("SELECT numero FROM os LIMIT 10")
                numeros_disponiveis = [row[0] for row in c.fetchall()]
                print(f"Números de OS disponíveis: {numeros_disponiveis}")
                mostrar_mensagem(self.root, "Erro", f"Não foi possível encontrar os dados da OS {numero_os}", "erro")
                conn.close()
                return
                
            print(f"OS encontrada no banco de dados: {resultado[0]}")
            cliente = resultado[1]  # índice 1 é o cliente
            
            # Verificar primeiro na nova estrutura de pastas
            pasta_os = self.get_pasta_os_cliente(numero_os, cliente)
            arquivo_excel = os.path.join(pasta_os, f"{numero_os}_visita_tecnica.xlsx")
            
            # Verificar se o arquivo existe na nova estrutura
            if os.path.exists(arquivo_excel):
                # Se o arquivo já existe, apenas abre sem recriar
                self.abrir_arquivo(arquivo_excel)
                mostrar_mensagem(self.root, "Sucesso", f"Documento existente aberto em:\n{arquivo_excel}", "sucesso")
                return
                
            # Verificar se o arquivo já existe no caminho antigo
            caminho_arquivo = self.get_caminho_arquivo('visita_tecnica', numero_os)
            if caminho_arquivo and os.path.exists(caminho_arquivo):
                # Se o arquivo já existe no local antigo, copiar para a nova estrutura
                os.makedirs(os.path.dirname(arquivo_excel), exist_ok=True)
                shutil.copy2(caminho_arquivo, arquivo_excel)
                
                # Abrir o arquivo da nova localização
                self.abrir_arquivo(arquivo_excel)
                mostrar_mensagem(self.root, "Sucesso", f"Documento existente copiado e aberto em:\n{arquivo_excel}", "sucesso")
                return
                
            # Se não existe, continua com a criação do documento
            # Gerar o documento usando a função existente - passar o número exato encontrado no banco de dados
            print(f"Gerando novo documento de OS Visita Técnica...")
            numero_os_para_criar = resultado[0]  # Usar o número exato como está no banco
            caminho_arquivo = criar_os_visita_tecnica(numero_os_para_criar)
            
            if caminho_arquivo:
                print(f"Documento gerado com sucesso em: {caminho_arquivo}")
                # Copiar para a nova estrutura
                os.makedirs(os.path.dirname(arquivo_excel), exist_ok=True)
                shutil.copy2(caminho_arquivo, arquivo_excel)
                
                # Abrir o documento com o aplicativo padrão (já da nova localização)
                self.abrir_arquivo(arquivo_excel)
                
                mostrar_mensagem(self.root, "Sucesso", f"OS de Visita Técnica gerada com sucesso em:\n{arquivo_excel}", "sucesso")
            else:
                print("Falha ao gerar o documento de OS Visita Técnica.")
                mostrar_mensagem(self.root, "Erro", "Não foi possível gerar a OS de Visita Técnica.", "erro")
        except Exception as e:
            print(f"Erro detalhado ao gerar OS de Visita Técnica: {str(e)}")
            import traceback
            traceback.print_exc()
            mostrar_mensagem(self.root, "Erro", f"Erro ao gerar OS de Visita Técnica: {str(e)}", "erro")

    def _configure_main_frame(self, event):
        """Atualiza a área de rolagem quando o frame muda de tamanho"""
        self.main_canvas.configure(scrollregion=self.main_canvas.bbox("all"))
    
    def _configure_main_canvas(self, event):
        """Ajusta a largura da janela interior quando o canvas é redimensionado"""
        self.main_canvas.itemconfig(self.frame_window, width=event.width)
    
    def _configure_mousewheel_scrolling(self):
        """Configura eventos de rolagem do mouse para permitir usar o scroll do mouse"""
        def _on_mousewheel(event):
            # Apenas rolar se o cursor do mouse estiver sobre o canvas principal
            widget_under_cursor = event.widget.winfo_containing(event.x_root, event.y_root)
            if widget_under_cursor and (widget_under_cursor == self.main_canvas or 
                                        widget_under_cursor.master == self.main_canvas or 
                                        widget_under_cursor.master == self.main_frame):
                # Linux: evento Button-4 (rolar para cima) e Button-5 (rolar para baixo)
                if event.num == 4:
                    self.main_canvas.yview_scroll(-1, "units")
                elif event.num == 5:
                    self.main_canvas.yview_scroll(1, "units")
                # Windows/macOS: evento MouseWheel com delta
                elif hasattr(event, 'delta'):
                    self.main_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")
        
        # Vincular eventos específicos por plataforma
        # Windows/macOS
        self.root.bind("<MouseWheel>", _on_mousewheel)
        # Linux
        self.root.bind("<Button-4>", _on_mousewheel)
        self.root.bind("<Button-5>", _on_mousewheel)

    def verificar_licenca_sistema(self):
        """Verifica se a licença do sistema é válida e inicia o processo de ativação se necessário."""
        try:
            # Tentar carregar licença
            codigo_licenca = carregar_licenca()
            
            # Se não existir código de licença
            if not codigo_licenca:
                return self.iniciar_ativacao_licenca()
            
            # Verificar licença existente
            resultado = verificar_licenca(codigo_licenca)
            
            # Se licença não for válida
            if not resultado["valida"]:
                # Se expirada, mostrar mensagem específica
                if "expirada" in resultado["mensagem"].lower():
                    mensagem = f"Sua licença expirou em {resultado.get('data_expiracao', 'data não especificada')}.\n"
                    mensagem += "Por favor, entre em contato para renovar sua licença."
                    messagebox.showerror("Licença Expirada", mensagem)
                else:
                    messagebox.showerror("Licença Inválida", resultado["mensagem"])
                
                return self.iniciar_ativacao_licenca()
            
            # Verificar se está próximo da expiração (menos de 30 dias)
            if resultado.get("dias_restantes", 0) < 30:
                mensagem = f"Sua licença expirará em {resultado['dias_restantes']} dias.\n"
                mensagem += "Entre em contato para renovar sua licença antes que expire."
                messagebox.showwarning("Licença Próxima de Expirar", mensagem)
            
            return True
            
        except Exception as e:
            messagebox.showerror("Erro de Licença", f"Erro ao verificar licença: {str(e)}")
            return self.iniciar_ativacao_licenca()
    
    def iniciar_ativacao_licenca(self):
        """Inicia o processo de ativação de licença."""
        # Criar janela temporária se self.root ainda não existir
        if not hasattr(self, 'root') or not self.root:
            root_temp = ttk.Window(themename="solar")
            root_temp.withdraw()  # Ocultar janela temporária
        else:
            root_temp = self.root
        
        # Função de callback para quando a licença é ativada
        def licenca_ativada(resultado):
            nonlocal continuar_inicializacao
            if resultado["valida"]:
                continuar_inicializacao = True
                if not hasattr(self, 'root') or not self.root:
                    root_temp.destroy()
        
        # Mostrar janela de ativação
        continuar_inicializacao = False
        ativador = AtivadorLicenca(root_temp, licenca_ativada)
        
        # Se estamos usando uma janela temporária, iniciar mainloop
        if not hasattr(self, 'root') or not self.root:
            root_temp.mainloop()
        
        return continuar_inicializacao

    def get_caminho_arquivo(self, tipo_os, numero_os):
        """
        Obtém o caminho onde um arquivo de OS deve estar armazenado, 
        baseado no tipo de OS e número da OS.
        
        Args:
            tipo_os (str): Tipo de OS (ziehm, andamento, interna, etc)
            numero_os (str): Número da OS
        
        Returns:
            str: Caminho completo do arquivo ou None se não puder determinar
        """
        # Formatação padronizada do número da OS
        if not numero_os.startswith("OS "):
            numeros = re.findall(r'\d+', numero_os)
            if numeros:
                numero_completo = ''.join(numeros)
                if len(numero_completo) >= 5:
                    ano = numero_completo[:2]
                    sequencia = numero_completo[2:].zfill(3)  # Garantir 3 dígitos
                    numero_os_formatado = f"OS {ano} {sequencia}"
                else:
                    ano = datetime.now().strftime("%y")
                    sequencia = numero_completo.zfill(3)  # Garantir 3 dígitos
                    numero_os_formatado = f"OS {ano} {sequencia}"
            else:
                ano = datetime.now().strftime("%y")
                numero_os_formatado = f"OS {ano} 001"
        else:
            numero_os_formatado = numero_os
        
        # Criar string limpa para uso em nome de arquivo
        numero_os_limpo = numero_os_formatado.replace(" ", "_").replace("/", "-").replace("-", "_")
        
        # Definir pasta principal para documentos técnicos
        pasta_base = os.path.join(os.path.dirname(os.path.abspath(__file__)), 
                                "MORACA", "MORACA1", "DOCUMENTOS", "tecnico")
        
        # Definir pasta de destino e padrão de nome do arquivo baseado no tipo de OS
        if tipo_os.lower() == 'ziehm':
            pasta_especifica = os.path.join(pasta_base, "andamento")
            padrao_arquivo = f"Ziehm_OS_{numero_os_limpo}"
        elif tipo_os.lower() in ['andamento', 'interna']:
            pasta_especifica = os.path.join(pasta_base, "andamento")
            padrao_arquivo = f"OS_{numero_os_limpo}"
        elif tipo_os.lower() == 'preventiva':
            pasta_especifica = os.path.join(pasta_base, "preventivas")
            padrao_arquivo = f"OS_Preventiva_{numero_os_limpo}"
        elif tipo_os.lower() == 'visita':
            pasta_especifica = os.path.join(pasta_base, "visitas")
            padrao_arquivo = f"OS_Visita_{numero_os_limpo}"
        elif tipo_os.lower() == 'visita_corretiva':
            pasta_especifica = os.path.join(pasta_base, "visitas")
            padrao_arquivo = f"OS_Visita_Corretiva_{numero_os_limpo}"
        elif tipo_os.lower() == 'visita_tecnica':
            pasta_especifica = os.path.join(pasta_base, "visitas")
            padrao_arquivo = f"OS_Visita_Tecnica_{numero_os_limpo}"
        else:
            # Para outros tipos de OS que ainda não têm um padrão definido
            return None
            
        # Verificar se a pasta existe, se não, criar
        if not os.path.exists(pasta_especifica):
            try:
                os.makedirs(pasta_especifica)
            except:
                pass
            
        # Tentar encontrar arquivo existente com o padrão
        arquivo_encontrado = self.encontrar_arquivo_por_padrao(pasta_especifica, padrao_arquivo)
        if arquivo_encontrado:
            return arquivo_encontrado
                
        # Se não encontrou, sugerir o nome padrão para criação
        return os.path.join(pasta_especifica, f"{padrao_arquivo}.xlsx")
    
    def encontrar_arquivo_por_padrao(self, pasta, padrao_arquivo):
        """
        Busca por arquivos que correspondam a um padrão em uma pasta.
        
        Args:
            pasta (str): Pasta onde buscar
            padrao_arquivo (str): Padrão de nome de arquivo a buscar
            
        Returns:
            str: Caminho completo do arquivo encontrado ou None
        """
        if not os.path.exists(pasta):
            return None
            
        # Listar todos os arquivos na pasta
        try:
            arquivos = os.listdir(pasta)
            for arquivo in arquivos:
                # Verificar se o arquivo corresponde ao padrão (ignorando extensão)
                if arquivo.startswith(padrao_arquivo):
                    return os.path.join(pasta, arquivo)
        except Exception as e:
            print(f"Erro ao listar arquivos: {str(e)}")
            
        return None

if __name__ == "__main__":
    # Iniciar a aplicação com verificação de licença
    app = MoracaOS()
    
    # Continuar apenas se a aplicação foi inicializada corretamente 
    # (o que significa que a licença é válida)
    if hasattr(app, 'root') and app.root:
        app.root.mainloop()
    else:
        sys.exit(0) 